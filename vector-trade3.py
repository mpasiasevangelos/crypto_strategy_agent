import os
import warnings
from dotenv import load_dotenv

os.environ['KMP_DUPLICATE_LIB_OK'] = 'True'
warnings.filterwarnings("ignore")

load_dotenv()

from langchain_community.document_loaders import PyMuPDFLoader
from docx import Document as DocxDocument
from ebooklib import epub
from langchain_core.documents import Document

# Function to load DOC/DOCX files
def load_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text, metadata={"source": file_path})]

# Function to load EPUB files
def load_epub(file_path):
    book = epub.read_epub(file_path)
    text = ""
    for item in book.items:
        if item.get_type() == epub.EpubHtml:
            text += item.content.decode("utf-8")
    return [Document(page_content=text, metadata={"source": file_path})]

# Collect all supported files
supported_extensions = [".pdf", ".docx", ".epub"]
files = []
for root, dirs, files_in_dir in os.walk("rag-dataset/crypto"):
    for file in files_in_dir:
        if any(file.endswith(ext) for ext in supported_extensions):
            files.append(os.path.join(root, file))

print("Supported files found:")
print(files)

# Load documents
docs = []
file_chunks = {}  # To store the chunk count for each file
for file in files:
    if file.endswith(".pdf"):
        print('\nLoading PDF:', file)
        loader = PyMuPDFLoader(file)
        temp_docs = loader.load()
    elif file.endswith(".docx"):
        print('\nLoading DOCX:', file)
        temp_docs = load_docx(file)
    elif file.endswith(".epub"):
        print('\nLoading EPUB:', file)
        temp_docs = load_epub(file)
    else:
        continue
    docs.extend(temp_docs)
    file_chunks[file] = len(temp_docs)  # Record the number of documents loaded for this file

print(f"Loaded {len(docs)} documents.")

# Split documents into chunks
from langchain_text_splitters import RecursiveCharacterTextSplitter

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
chunks = []
for file, doc_count in file_chunks.items():
    # Filter out documents from the current file
    file_docs = [doc for doc in docs if doc.metadata["source"] == file]
    file_chunks[file] = text_splitter.split_documents(file_docs)
    print(f"File '{file}' generated {len(file_chunks[file])} chunks.")  # Log chunk count for each file
    chunks.extend(file_chunks[file])

print(f"Split into {len(chunks)} chunks.")

# Embed chunks and create a vector store
from langchain_ollama import OllamaEmbeddings
import faiss
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore

embeddings = OllamaEmbeddings(
    model="nomic-embed-text",
    base_url='http://localhost:11434'
)

# Create a FAISS index
vector = embeddings.embed_query(chunks[0].page_content)
index = faiss.IndexFlatIP(len(vector))

vector_store = FAISS(
    embedding_function=embeddings,
    index=index,
    docstore=InMemoryDocstore(),
    index_to_docstore_id={}
)

# Add chunks to the vector store
ids = vector_store.add_documents(documents=chunks)
print(f"Added {len(ids)} chunks to the vector store.")

# Save the vector store locally
db_name = "trade"
vector_store.save_local(db_name)
print(f"Vector store saved to '{db_name}'.")
